""" Widgets for showing the document structure.
"""
import cProfile
import logging
import os.path
import pstats

from PySide6 import QtCore, QtWidgets
from PySide6.QtCore import Qt, Signal

from docx import Document as createDocument
from docx.document import Document
from docx.text.paragraph import Paragraph

from wuw.be.info import PROFILING
from wuw.gui.togglecolumn import ToggleColumnTableView
from wuw.utils.cls import check_type

logger = logging.getLogger(__name__)

# Header widths for different column types
_HW_STR = 100 # Small string label
_HW_BOOL = 80
_HW_INT = 100

_ALIGN_STRING = int(Qt.AlignVCenter | Qt.AlignLeft)
_ALIGN_NUMBER = int(Qt.AlignVCenter | Qt.AlignRight)
_ALIGN_BOOLEAN = int(Qt.AlignVCenter | Qt.AlignHCenter)

ALL_ITEMS_STR = "All"



class DocumentTableModel(QtCore.QAbstractTableModel):
    """ A table model that maps the document to a table.
    """
    HEADERS = ('Text', 'Justification', 'Alignment',
               '1st line indent', 'Left indent', 'Right indent')

    HEADER_TOOL_TIPS = HEADERS

    (COL_NAME, COL_JUSTIFICATION, COL_ALIGNMENT,
     COL_1ST_LINE_INDENT, COL_LEFT_INDENT, COL_RIGHT_INDENT) = range(len(HEADERS))

    DEFAULT_WIDTHS = [400, _HW_STR, _HW_STR,
                      _HW_INT, _HW_INT, _HW_INT]

    SORT_ROLE = Qt.UserRole

    def __init__(self, document: Document=None, **kwargs):
        """ Constructor

            :param CmLib cmLib: the underlying color library
            :param QWidget parent: Qt parent widget
        """
        super().__init__(**kwargs)
        check_type(document, Document, allow_none=True)

        assert len(self.HEADERS) == len(self.DEFAULT_WIDTHS), "sanity check failed."
        assert len(self.HEADERS) == len(self.HEADER_TOOL_TIPS), "sanity check failed."

        self._document = createDocument() if document is None else document

        # Check mark for boolean columns
        #   ✓ Check mark Unicode: U+2713
        #   ✔︎ Heavy check mark Unicode: U+2714
        self.checkmarkChar = '✓︎'

        if PROFILING:
            # Profiler that measures the drawing of the inspectors.
            self._profFileName = "wuw.prof"

            logger.warning("Profiling is on for {}. This may cost a bit of CPU time.")
            self._profiler = cProfile.Profile()

    def finalize(self):
        if PROFILING:
            logger.info("Saving profiling information to {}"
                        .format(os.path.abspath(self._profFileName)))
            self._profiler.create_stats()

            profStats = pstats.Stats(self._profiler)
            profStats.dump_stats(self._profFileName)

    @property
    def document(self) -> Document:
        """ Returns the underlying document
        """
        return self._document

    @document.setter
    def document(self, doc: Document) -> None:
        """ Sets the document as data in the model

            Resets the model.
        """
        self.beginResetModel()
        try:
            self._document = createDocument() if doc is None else doc
        finally:
            self.endResetModel()


    @property
    def paragraphs(self):
        """ Returns the paragraphs of the underlying document
        """
        return self._document.paragraphs



    def rowCount(self, _parent=None):
        """ Returns the number of rows.
        """
        return len(self.paragraphs)


    def columnCount(self, _parent=None):
        """ Returns the number of columns.
        """
        return len(self.HEADERS)


    def _posFromIndex(self, index):
        """ Returns the (row, col) given the index

            Returns None if the index is invalid or row or column are negative or larger than the
            number of rows/cols
        """
        if not index.isValid():
            return None

        row, col = index.row(), index.column()

        if col < 0 or col >= self.columnCount():
            return None

        if row < 0 or row >= self.rowCount():
            return None

        return row, col


    def flags(self, index):
        """ Returns the item flags for the given index
        """
        pos = self._posFromIndex(index)
        if pos is None:
            return None
        else:
            row, col = pos

        result = Qt.ItemIsEnabled | Qt.ItemIsSelectable

        # if col == self.COL_FAV:
        #     result = result | Qt.ItemIsUserCheckable | Qt.ItemIsEditable
        return result


    def _boolToData(self, value, role=Qt.DisplayRole):
        """ If role is the display roel it shows a checkmark if value is True, the empty string
            otherwise. For other roles it just returns the boolean value
        """
        if role == Qt.DisplayRole:
            return self.checkmarkChar if value else ""
        else:
            return bool(value) # convert to bool just in case


    def data(self, index, role=Qt.DisplayRole):
        """ Returns the data stored under the given role for the item referred to by the index.
        """
        pos = self._posFromIndex(index)
        if pos is None:
            return None
        else:
            row, col = pos

        match role:
            case Qt.DisplayRole | self.SORT_ROLE | Qt.ToolTipRole:
                try:
                    if PROFILING:
                        self._profiler.enable()

                    return "fna"
                    paragraph = self.paragraphs[row]

                    match col:
                        case self.COL_NAME:
                            return paragraph.text
                        case self.COL_JUSTIFICATION:
                            # TODO: Is there a difference with paragraph_format.alignment?
                            return str(paragraph.alignment)
                        case self.COL_ALIGNMENT:
                            return str(paragraph.paragraph_format.alignment)
                        case self.COL_1ST_LINE_INDENT:
                            return str(paragraph.paragraph_format.first_line_indent)
                        case self.COL_LEFT_INDENT:
                            return str(paragraph.paragraph_format.left_indent)
                        case self.COL_RIGHT_INDENT:
                            return str(paragraph.paragraph_format.right_indent)
                        case _ :
                            raise AssertionError("Unexpected column: {}".format(col))

                finally:
                    if PROFILING:
                        self._profiler.disable()

            case Qt.TextAlignmentRole:
                return _ALIGN_STRING

        return None


    def headerData(self, section, orientation, role=Qt.DisplayRole):
        """ Header data given a orientation.

            :param section: row or column number, depending on orientation
            :param orientation: Qt.Horizontal or Qt.Vertical
        """
        if role not in (Qt.DisplayRole, Qt.ToolTipRole):
            return None

        if orientation == Qt.Horizontal:
            match role:
                case Qt.DisplayRole:
                    return self.HEADERS[section]
                case Qt.ToolTipRole: # doesn't seem to work (tried only on OS-X)
                    return self.HEADER_TOOL_TIPS[section]
                case _:
                    assert False, "Unexpected role: {}".format(role)
        else:
            return str(section)


    def getParagraphByIndex(self, index):
        """ Returns a color map at row of the given index.

            Returns None if the index is not valid
        """
        if not index.isValid():
            return None
        else:
            return self.paragraphs[index.row()]



class DocumentTableViewer(ToggleColumnTableView):

    sigParagraphHighlighted = Signal(object)  # Optional paragraph

    def __init__(self, model: DocumentTableModel=None, parent=None):
        """ Constructor

            :param DocumentTableModel model: the item model
        """
        super().__init__()

        check_type(model, DocumentTableModel)
        self._sourceModel = model
        self.setModel(self._sourceModel)

        self.setWordWrap(False)
        self.setSortingEnabled(False)
        self.setShowGrid(False)
        self.setCornerButtonEnabled(True)
        self.setAlternatingRowColors(True)
        self.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows)
        self.setSelectionMode(QtWidgets.QAbstractItemView.SingleSelection)

        self.verticalHeader().show()
        horHeader = self.horizontalHeader()
        horHeader.setSectionsMovable(True)
        horHeader.setSectionResizeMode(QtWidgets.QHeaderView.ResizeMode.Interactive)
        # This unfortunately does not work intuitively when resizing headers.
        #horHeader.setSectionResizeMode(0, QtWidgets.QHeaderView.ResizeMode.Stretch)
        horHeader.setStretchLastSection(True)

        for col, width in enumerate(DocumentTableModel.DEFAULT_WIDTHS):
            horHeader.resizeSection(col, width)

        # Make the 'name' color wider because of the legend bar.
        horHeader.resizeSection(
            DocumentTableModel.COL_NAME,
            DocumentTableModel.DEFAULT_WIDTHS[DocumentTableModel.COL_NAME])

        headerNames = DocumentTableModel.HEADERS
        enabled = dict((name, True) for name in headerNames)
        enabled[headerNames[DocumentTableModel.COL_NAME]] = False # Cannot be unchecked
        checked = dict((name, True) for name in headerNames)
        self.addHeaderContextMenu(checked=checked, enabled=enabled, checkable={})

        self.setContextMenuPolicy(Qt.DefaultContextMenu) # will call contextMenuEvent

        self._selectionModel = self.selectionModel()
        self._selectionModel.currentChanged.connect(self._onCurrentChanged)


    def _onCurrentChanged(self, curIdx, _prevIdx):
        """ Emits sigColorMapSelected if a valid row has been selected
        """
        if not curIdx.isValid():
            logger.info("NONE selected")
            paragraph = None
        else:
            assert curIdx.isValid(), "Current index not valid"

            row = curIdx.row()
            paragraph = self._sourceModel.document.paragraphs[row]

        logger.debug("Emitting sigParagraphHighlighted: {}".format(paragraph))
        self.sigParagraphHighlighted.emit(paragraph)


    def getCurrentColorMap(self):
        """ Gets the current colorMap

            Returns None if None selected.
        """
        return self._sourceModel.getParagraphByIndex(self.currentIndex())


    def scrollToCurrent(self):
        """ Scroll to the currently selected color map.
        """
        curIdx = self.currentIndex()
        logger.debug("scrollToCurrent: {} (isValid)".format(curIdx, curIdx.isValid()))
        if curIdx.isValid():
            self.scrollTo(self.currentIndex())

